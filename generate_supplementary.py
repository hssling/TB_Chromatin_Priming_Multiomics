"""
Generate Supplementary Materials for IJTLD submission
Extended methods, additional tables, and supporting data
"""

import os
import csv
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "4_results"
TABLES_DIR = RESULTS_DIR / "tables"
FIGURES_DIR = RESULTS_DIR / "figures"
OUTPUT_DIR = RESULTS_DIR / "submission"
OUTPUT_DIR.mkdir(exist_ok=True)

def read_csv(path, max_rows=None):
    if not path.exists():
        return []
    with open(path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        if max_rows:
            return [next(reader) for _ in range(max_rows) if reader]
        return list(reader)

def add_table_from_csv(doc, csv_path, title, max_rows=50):
    """Add a table from CSV file"""
    data = read_csv(csv_path, max_rows)
    if not data:
        doc.add_paragraph(f"[Data not available: {csv_path.name}]")
        return
    
    if title:
        t = doc.add_paragraph(title)
        t.runs[0].bold = True
    
    headers = list(data[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in data[:max_rows]:
        row = table.add_row()
        for i, h in enumerate(headers):
            value = row_data.get(h, '')
            # Format numeric values
            try:
                if '.' in str(value):
                    value = f"{float(value):.4f}"
            except:
                pass
            row.cells[i].text = str(value)[:50]  # Truncate long values
    
    doc.add_paragraph()

def main():
    print("Generating Supplementary Materials...")
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ===== TITLE =====
    title = doc.add_heading("Supplementary Materials", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(
        "Chromatin Priming Index: A Novel Quantitative Framework for Integrating "
        "Single-Cell Transcriptomic and Epigenomic Data in Tuberculosis"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph("Siddalingaiah H S").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading("Contents", level=1)
    toc = [
        "Supplementary Methods",
        "  - S1. Data preprocessing parameters",
        "  - S2. Cell type annotation markers",
        "  - S3. CPI calculation algorithm",
        "  - S4. Master regulator scoring",
        "Supplementary Tables",
        "  - Table S1. Complete DEG list (BAL dataset)",
        "  - Table S2. Complete DEG list (DTB PBMC dataset)",
        "  - Table S3. All transcription factors analyzed",
        "  - Table S4. Peak-gene link statistics",
        "Supplementary Figures",
        "  - Figure S1. Volcano plots by cell type",
        "  - Figure S2. Trained immunity module scores",
        "  - Figure S3. Metabolic signature accessibility",
        "Code Availability",
    ]
    for item in toc:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ===== SUPPLEMENTARY METHODS =====
    doc.add_heading("Supplementary Methods", level=1)
    
    doc.add_heading("S1. Data preprocessing parameters", level=2)
    doc.add_paragraph(
        "Single-cell RNA sequencing data were processed using the following parameters:"
    )
    
    params = doc.add_table(rows=1, cols=3)
    params.style = 'Table Grid'
    params.rows[0].cells[0].text = 'Parameter'
    params.rows[0].cells[1].text = 'Value'
    params.rows[0].cells[2].text = 'Description'
    
    param_data = [
        ("min.cells", "3", "Minimum cells per gene"),
        ("min.features", "200", "Minimum genes per cell"),
        ("nfeatures", "3000", "Variable features for PCA"),
        ("npcs", "30", "Principal components"),
        ("dims", "1:20", "Dimensions for clustering"),
        ("resolution", "0.8", "Clustering resolution"),
        ("n.neighbors", "30", "UMAP neighbors"),
    ]
    for p in param_data:
        row = params.add_row()
        row.cells[0].text = p[0]
        row.cells[1].text = p[1]
        row.cells[2].text = p[2]
    
    doc.add_paragraph()
    
    doc.add_heading("S2. Cell type annotation markers", level=2)
    doc.add_paragraph(
        "Cell types were annotated using canonical marker gene expression. Module scores "
        "were calculated using Seurat's AddModuleScore function with nbin=24."
    )
    
    markers = doc.add_table(rows=1, cols=2)
    markers.style = 'Table Grid'
    markers.rows[0].cells[0].text = 'Cell Type'
    markers.rows[0].cells[1].text = 'Marker Genes'
    
    marker_data = [
        ("T cell", "CD3D, CD3E, CD4, CD8A"),
        ("NK cell", "GNLY, NKG7, NCAM1"),
        ("B cell", "CD79A, MS4A1, CD19"),
        ("Monocyte", "CD14, LYZ, S100A8, S100A9"),
        ("Dendritic cell", "CD1C, FCER1A, CLEC4C"),
        ("Macrophage (alveolar)", "MARCO, FABP4, MCEMP1"),
        ("Macrophage (interstitial)", "CD68, CD163, MRC1"),
    ]
    for m in marker_data:
        row = markers.add_row()
        row.cells[0].text = m[0]
        row.cells[1].text = m[1]
    
    doc.add_paragraph()
    
    doc.add_heading("S3. CPI calculation algorithm", level=2)
    doc.add_paragraph("The Chromatin Priming Index was calculated as follows:")
    
    algo = doc.add_paragraph()
    algo.add_run("Algorithm 1: CPI Calculation").bold = True
    
    code_block = """
Input: DEG_list (differentially expressed genes)
       peak_links (peak-gene associations from reference)
Output: CPI (chromatin priming index, 0-1)

1. Filter DEGs: sig_DEGs = DEGs with adj.P < 0.05
2. Get linked genes: linked_genes = unique genes in peak_links
3. Calculate intersection: primed_DEGs = sig_DEGs ∩ linked_genes
4. Compute CPI: CPI = |primed_DEGs| / |sig_DEGs|
5. Return CPI
"""
    for line in code_block.strip().split('\n'):
        doc.add_paragraph(line, style='Quote')
    
    doc.add_heading("S4. Master regulator scoring", level=2)
    doc.add_paragraph(
        "Transcription factors were ranked using a composite score incorporating multiple factors:"
    )
    
    formula = doc.add_paragraph()
    formula.add_run("Score = N_cell_types × |mean(log2FC)| × priming_bonus").italic = True
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "where N_cell_types is the number of cell types in which the TF is significantly "
        "differentially expressed, mean(log2FC) is the average fold-change across cell types, "
        "and priming_bonus = 1.5 if the TF has chromatin accessibility support, else 1.0."
    )
    
    doc.add_page_break()
    
    # ===== SUPPLEMENTARY TABLES =====
    doc.add_heading("Supplementary Tables", level=1)
    
    # Table S1: Top DEGs BAL
    doc.add_heading("Table S1. Top 50 Differentially Expressed Genes (BAL Dataset)", level=2)
    add_table_from_csv(doc, TABLES_DIR / "DEG_all_celltypes.csv", "", max_rows=50)
    
    # Table S2: Top DEGs DTB
    doc.add_heading("Table S2. Top 50 Differentially Expressed Genes (DTB PBMC Dataset)", level=2)
    add_table_from_csv(doc, TABLES_DIR / "DEG_GSE287288_DTB.csv", "", max_rows=50)
    
    # Table S3: All TFs
    doc.add_heading("Table S3. All Transcription Factors Analyzed", level=2)
    add_table_from_csv(doc, TABLES_DIR / "Table_TF_master_regulators.csv", "", max_rows=60)
    
    # Table S4: Peak-gene link stats
    doc.add_heading("Table S4. Peak-Gene Link Statistics", level=2)
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.rows[0].cells[0].text = 'Metric'
    stats_table.rows[0].cells[1].text = 'Value'
    
    stats_data = [
        ("Total peak-gene links", "1,283,042"),
        ("Number of cells in reference", "12,012"),
        ("Number of peaks", "111,857"),
        ("Unique genes with links", "~15,000"),
        ("Average links per gene", "~85"),
        ("Reference dataset", "10x Genomics PBMC Multiome"),
        ("Link correlation threshold", "|r| > 0.2"),
    ]
    for s in stats_data:
        row = stats_table.add_row()
        row.cells[0].text = s[0]
        row.cells[1].text = s[1]
    
    doc.add_page_break()
    
    # ===== SUPPLEMENTARY FIGURES =====
    doc.add_heading("Supplementary Figures", level=1)
    
    # Figure S1
    fig_path = FIGURES_DIR / "Fig3_volcano_DEG.png"
    if fig_path.exists():
        doc.add_heading("Figure S1. Volcano Plot of Differentially Expressed Genes", level=2)
        doc.add_picture(str(fig_path), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Volcano plot showing differentially expressed genes in BAL macrophages. "
            "Red points indicate significantly upregulated genes (log2FC > 0.5, adj.P < 0.05), "
            "blue points indicate significantly downregulated genes."
        ).italic = True
    
    # Figure S2
    fig_path = FIGURES_DIR / "Fig5A_trained_immunity_score.png"
    if fig_path.exists():
        doc.add_heading("Figure S2. Trained Immunity Module Scores", level=2)
        doc.add_picture(str(fig_path), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "UMAP visualization colored by trained immunity module score. Higher scores "
            "indicate enrichment of trained immunity-associated genes."
        ).italic = True
    
    # Figure S3
    fig_path = FIGURES_DIR / "Fig6_metabolic_accessibility.png"
    if fig_path.exists():
        doc.add_heading("Figure S3. Metabolic Gene Chromatin Accessibility", level=2)
        doc.add_picture(str(fig_path), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Comparison of chromatin accessibility scores for different gene categories. "
            "Boxplots show median and interquartile range."
        ).italic = True
    
    # Figure S4
    fig_path = FIGURES_DIR / "Fig4B_priming_categories.png"
    if fig_path.exists():
        doc.add_heading("Figure S4. Chromatin Priming Categories", level=2)
        doc.add_picture(str(fig_path), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Distribution of genes across chromatin priming categories: primed (accessible chromatin), "
            "non-primed (no detected accessibility), and background."
        ).italic = True
    
    doc.add_page_break()
    
    # ===== CODE AVAILABILITY =====
    doc.add_heading("Code Availability", level=1)
    
    doc.add_paragraph(
        "All analysis code is available at the following GitHub repository:"
    )
    
    doc.add_paragraph("https://github.com/hssling/TB_Chromatin_Priming_Multiomics")
    
    doc.add_paragraph("The repository includes:")
    code_items = [
        "R scripts for single-cell RNA-seq analysis (Seurat v5)",
        "Python scripts for manuscript generation",
        "Configuration files for reproducibility",
        "Sample data and expected outputs",
        "GitHub Actions CI/CD workflow",
    ]
    for item in code_items:
        doc.add_paragraph(f"• {item}")
    
    doc.add_heading("Software versions", level=2)
    sw_table = doc.add_table(rows=1, cols=2)
    sw_table.style = 'Table Grid'
    sw_table.rows[0].cells[0].text = 'Software'
    sw_table.rows[0].cells[1].text = 'Version'
    
    software = [
        ("R", "4.5.1"),
        ("Seurat", "5.0+"),
        ("data.table", "1.14+"),
        ("ggplot2", "3.4+"),
        ("Python", "3.11"),
        ("python-docx", "0.8.11"),
    ]
    for sw in software:
        row = sw_table.add_row()
        row.cells[0].text = sw[0]
        row.cells[1].text = sw[1]
    
    # Save
    output_path = OUTPUT_DIR / "Supplementary_Materials_IJTLD.docx"
    doc.save(output_path)
    print(f"Supplementary materials saved to: {output_path}")

if __name__ == "__main__":
    main()
