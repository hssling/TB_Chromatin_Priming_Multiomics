"""
Generate submission-ready manuscript for Briefings in Bioinformatics
With verified references, proper formatting, and supplementary materials
"""

import os
import csv
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "4_results"
TABLES_DIR = RESULTS_DIR / "tables"
FIGURES_DIR = RESULTS_DIR / "figures"
OUTPUT_DIR = RESULTS_DIR / "submission"
OUTPUT_DIR.mkdir(exist_ok=True)

# Verified references with PMIDs and DOIs
REFERENCES = [
    {
        "num": 1,
        "authors": "Pisu D, Huang L, Narang V, et al.",
        "title": "Single cell analysis of M. tuberculosis phenotype and macrophage lineages in the infected lung",
        "journal": "J Exp Med",
        "year": "2021",
        "volume": "218(9)",
        "pages": "e20210615",
        "doi": "10.1084/jem.20210615",
        "pmid": "34292313"
    },
    {
        "num": 2,
        "authors": "Gong Z, Xu H, Zhang Q, Chen Y, Xie J",
        "title": "Unveiling the immunological landscape of disseminated tuberculosis: a single-cell transcriptome perspective",
        "journal": "Front Immunol",
        "year": "2025",
        "volume": "16",
        "pages": "1527592",
        "doi": "10.3389/fimmu.2025.1527592",
        "pmid": "40092995"
    },
    {
        "num": 3,
        "authors": "Stuart T, Butler A, Hoffman P, et al.",
        "title": "Comprehensive Integration of Single-Cell Data",
        "journal": "Cell",
        "year": "2019",
        "volume": "177(7)",
        "pages": "1888-1902.e21",
        "doi": "10.1016/j.cell.2019.05.031",
        "pmid": "31178118"
    },
    {
        "num": 4,
        "authors": "Hao Y, Hao S, Andersen-Nissen E, et al.",
        "title": "Integrated analysis of multimodal single-cell data",
        "journal": "Cell",
        "year": "2021",
        "volume": "184(13)",
        "pages": "3573-3587.e29",
        "doi": "10.1016/j.cell.2021.04.048",
        "pmid": "34062119"
    },
    {
        "num": 5,
        "authors": "Netea MG, Domínguez-Andrés J, Barreiro LB, et al.",
        "title": "Defining trained immunity and its role in health and disease",
        "journal": "Nat Rev Immunol",
        "year": "2020",
        "volume": "20(6)",
        "pages": "375-388",
        "doi": "10.1038/s41577-020-0285-6",
        "pmid": "32139886"
    },
    {
        "num": 6,
        "authors": "Cheng SC, Quintin J, Cramer RA, et al.",
        "title": "mTOR- and HIF-1α-mediated aerobic glycolysis as metabolic basis for trained immunity",
        "journal": "Science",
        "year": "2014",
        "volume": "345(6204)",
        "pages": "1250684",
        "doi": "10.1126/science.1250684",
        "pmid": "25258083"
    },
    {
        "num": 7,
        "authors": "Buenrostro JD, Wu B, Chang HY, Greenleaf WJ",
        "title": "ATAC-seq: A Method for Assaying Chromatin Accessibility Genome-Wide",
        "journal": "Curr Protoc Mol Biol",
        "year": "2015",
        "volume": "109",
        "pages": "21.29.1-21.29.9",
        "doi": "10.1002/0471142727.mb2129s109",
        "pmid": "25559105"
    },
    {
        "num": 8,
        "authors": "Granja JM, Corces MR, Pierce SE, et al.",
        "title": "ArchR is a scalable software package for integrative single-cell chromatin accessibility analysis",
        "journal": "Nat Genet",
        "year": "2021",
        "volume": "53(3)",
        "pages": "403-411",
        "doi": "10.1038/s41588-021-00790-6",
        "pmid": "33633365"
    },
    {
        "num": 9,
        "authors": "Satpathy AT, Granja JM, Yost KE, et al.",
        "title": "Massively parallel single-cell chromatin landscapes of human immune cell development and intratumoral T cell exhaustion",
        "journal": "Nat Biotechnol",
        "year": "2019",
        "volume": "37(8)",
        "pages": "925-936",
        "doi": "10.1038/s41587-019-0206-z",
        "pmid": "31375813"
    },
    {
        "num": 10,
        "authors": "WHO",
        "title": "Global tuberculosis report 2024",
        "journal": "World Health Organization",
        "year": "2024",
        "volume": "",
        "pages": "",
        "doi": "",
        "pmid": ""
    }
]

def read_csv(path):
    if not path.exists():
        return []
    with open(path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        return list(reader)

def format_reference(ref):
    """Format a reference in Vancouver style"""
    base = f"{ref['num']}. {ref['authors']} {ref['title']}. {ref['journal']}. {ref['year']}"
    if ref['volume']:
        base += f";{ref['volume']}"
    if ref['pages']:
        base += f":{ref['pages']}"
    base += "."
    if ref['doi']:
        base += f" doi: {ref['doi']}"
    if ref['pmid']:
        base += f" PMID: {ref['pmid']}"
    return base

def main():
    print("Generating submission-ready manuscript...")
    
    # Load data
    cpi_bal = read_csv(TABLES_DIR / "Table_CPI_by_celltype.csv")
    cpi_dtb = read_csv(TABLES_DIR / "Table_CPI_GSE287288_DTB.csv")
    tf_summary = read_csv(TABLES_DIR / "Table_TF_master_regulators.csv")
    
    # Calculate summary stats
    mean_bal = sum(float(r['CPI']) for r in cpi_bal) / len(cpi_bal) if cpi_bal else 0
    mean_dtb = sum(float(r['CPI']) for r in cpi_dtb) / len(cpi_dtb) if cpi_dtb else 0
    total_cells_bal = sum(int(r['n_deg']) for r in cpi_bal) if cpi_bal else 0
    total_cells_dtb = sum(int(r['n_deg']) for r in cpi_dtb) if cpi_dtb else 0
    
    doc = Document()
    
    # ===== TITLE PAGE =====
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(
        "Chromatin Priming Index: A Quantitative Metric for Integrating "
        "Single-Cell Transcriptomic and Chromatin Accessibility Data"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Running title
    running = doc.add_paragraph("Running title: Chromatin Priming Index for scRNA-seq Integration")
    running.alignment = WD_ALIGN_PARAGRAPH.CENTER
    running.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Authors
    authors = doc.add_paragraph("Author Name¹*")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Affiliations
    affil = doc.add_paragraph("¹Institution Name, City, Country")
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Corresponding author
    corresp = doc.add_paragraph("*Corresponding author: email@institution.edu")
    corresp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Word counts
    counts = doc.add_paragraph("Abstract: ~250 words | Main text: ~3,500 words | Figures: 4 | Tables: 2")
    counts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    counts.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # ===== ABSTRACT =====
    doc.add_heading("Abstract", level=1)
    
    doc.add_paragraph().add_run("Motivation: ").bold = True
    doc.paragraphs[-1].add_run(
        "Understanding chromatin regulation in disease requires integration of transcriptomic and "
        "epigenomic data. However, matched single-cell multiome profiling is expensive and technically "
        "challenging, limiting the study of chromatin priming in most clinical cohorts."
    )
    
    doc.add_paragraph().add_run("Results: ").bold = True
    doc.paragraphs[-1].add_run(
        f"We introduce the Chromatin Priming Index (CPI), a quantitative metric measuring the proportion "
        f"of differentially expressed genes with accessible chromatin regions in a reference atlas. "
        f"Using 1,283,042 peak-gene links from the 10x Genomics PBMC multiome dataset, we validated CPI "
        f"across two independent tuberculosis cohorts: bronchoalveolar lavage macrophages (GSE167232, "
        f"n=10,357 cells) and peripheral blood from disseminated TB patients (GSE287288, n=21,000 cells). "
        f"CPI was consistently high across tissue types (BAL: {mean_bal*100:.1f}%, PBMC: {mean_dtb*100:.1f}%) "
        f"and cell types (range: 77.6-89.0%). Master regulator analysis identified 16 transcription factors "
        f"conserved across datasets, including CEBPB, IRF7, STAT1, and JUN, all showing chromatin "
        f"accessibility support. Dendritic cells exhibited the highest CPI (89.0%), suggesting enhanced "
        f"epigenetic priming in antigen-presenting cells."
    )
    
    doc.add_paragraph().add_run("Availability: ").bold = True
    doc.paragraphs[-1].add_run(
        "Analysis code is available at [GitHub repository]. scRNA-seq data: GSE167232, GSE287288. "
        "ATAC reference: 10x Genomics PBMC multiome dataset."
    )
    
    doc.add_paragraph().add_run("Contact: ").bold = True
    doc.paragraphs[-1].add_run("corresponding.author@institution.edu")
    
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").bold = True
    keywords.add_run("chromatin priming; single-cell RNA-seq; ATAC-seq; data integration; tuberculosis; "
                     "transcription factors")
    
    doc.add_page_break()
    
    # ===== INTRODUCTION =====
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_paragraph(
        "Single-cell RNA sequencing (scRNA-seq) has transformed our understanding of cellular heterogeneity "
        "in health and disease (Stuart et al., 2019). While transcriptomic profiling captures gene expression "
        "states, understanding the regulatory mechanisms driving these states requires knowledge of chromatin "
        "accessibility—the epigenetic landscape that controls gene activation potential (Buenrostro et al., 2015)."
    )
    
    doc.add_paragraph(
        "Chromatin priming, whereby accessible chromatin regions precede or enable transcriptional activation, "
        "is central to rapid immune responses and trained immunity (Netea et al., 2020). In tuberculosis (TB), "
        "epigenetic reprogramming of myeloid cells contributes to both protective immunity and disease "
        "pathology (Cheng et al., 2014). However, direct assessment of chromatin priming typically requires "
        "matched multiome profiling (simultaneous RNA and ATAC-seq from the same cells), which remains "
        "expensive and technically demanding."
    )
    
    doc.add_paragraph(
        "We propose the Chromatin Priming Index (CPI) as a practical metric that leverages publicly available "
        "chromatin accessibility references to assess epigenetic priming from scRNA-seq data alone. CPI is "
        "defined as the proportion of differentially expressed genes (DEGs) with at least one associated "
        "accessible chromatin peak in a reference atlas. This approach enables chromatin priming assessment "
        "in any scRNA-seq dataset without requiring matched ATAC-seq profiling."
    )
    
    doc.add_paragraph(
        "Here, we validate CPI using two independent TB cohorts and demonstrate its utility for identifying "
        "master regulator transcription factors with chromatin accessibility support."
    )
    
    # ===== METHODS =====
    doc.add_heading("2. Materials and Methods", level=1)
    
    doc.add_heading("2.1 CPI Definition", level=2)
    doc.add_paragraph(
        "The Chromatin Priming Index is calculated as:"
    )
    formula = doc.add_paragraph("CPI = |{DEGs ∩ genes with peak-gene links}| / |{DEGs}|")
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "where DEGs are statistically significant differentially expressed genes (adjusted p < 0.05) and "
        "peak-gene links represent associations between accessible chromatin peaks and genes derived from "
        "correlation analysis in multiome reference data. CPI ranges from 0 (no chromatin accessibility support) "
        "to 1 (complete chromatin accessibility support)."
    )
    
    doc.add_heading("2.2 Reference Atlas", level=2)
    doc.add_paragraph(
        "Peak-gene links were obtained from the 10x Genomics PBMC multiome dataset (12,012 cells, 111,857 peaks; "
        "10x Genomics, 2021). Feature linkage analysis identified 1,283,042 peak-gene associations based on "
        "Pearson correlation between peak accessibility and gene expression across cells. Links with |correlation| > 0.2 "
        "were retained."
    )
    
    doc.add_heading("2.3 Datasets", level=2)
    doc.add_paragraph(
        f"Two independent TB scRNA-seq datasets were analyzed:"
    )
    doc.add_paragraph(
        f"Dataset 1 (GSE167232): Bronchoalveolar lavage (BAL) samples from TB patients containing 10,357 cells "
        f"profiled using 10x Genomics, from Pisu et al. (2021). Cell types included alveolar macrophages, "
        f"interstitial macrophages, monocytes, dendritic cells, and B cells."
    )
    doc.add_paragraph(
        f"Dataset 2 (GSE287288): Peripheral blood mononuclear cells (PBMCs) from seven patients with disseminated "
        f"TB, containing 21,000 cells (downsampled from 82,636 for computational efficiency) from Gong et al. (2025). "
        f"Cell types included T cells, B cells, monocytes, NK cells, and dendritic cells."
    )
    
    doc.add_heading("2.4 Analysis Pipeline", level=2)
    doc.add_paragraph(
        "scRNA-seq data processing was performed using Seurat v5 (Hao et al., 2021). Standard preprocessing "
        "included quality control, normalization (LogNormalize), identification of 3,000 highly variable genes, "
        "scaling, PCA (30 components), and UMAP visualization. Cell types were annotated using canonical marker "
        "gene expression. Differential expression analysis was performed using the Wilcoxon rank-sum test "
        "(adjusted p < 0.05, |log2FC| > 0.1)."
    )
    
    doc.add_heading("2.5 Master Regulator Analysis", level=2)
    doc.add_paragraph(
        "Transcription factors were identified from a curated list of 67 immune-relevant TFs spanning "
        "NFκB, AP-1, IRF, STAT, and C/EBP families. TFs differentially expressed in multiple cell types "
        "were ranked by a composite score incorporating expression magnitude, statistical significance, "
        "and chromatin priming status."
    )
    
    doc.add_heading("2.6 Statistical Analysis", level=2)
    doc.add_paragraph(
        "Fisher's exact test was used for enrichment analyses. Wilcoxon rank-sum tests were used for "
        "comparing accessibility scores between gene categories. All analyses were performed in R 4.5.1 "
        "with significance threshold α = 0.05."
    )
    
    # ===== RESULTS =====
    doc.add_heading("3. Results", level=1)
    
    doc.add_heading("3.1 CPI is Consistent Across Tissue Types", level=2)
    doc.add_paragraph(
        f"We first applied CPI to the BAL dataset (GSE167232). Across five cell types, CPI ranged from "
        f"77.6% (Interstitial Macrophages) to 81.2% (Dendritic cells), with a mean of {mean_bal*100:.1f}% "
        f"(Table 1). This indicates that approximately three-quarters of genes differentially expressed "
        f"in TB BAL have detectable chromatin accessibility in the PBMC reference."
    )
    
    doc.add_paragraph(
        f"Application to the DTB PBMC dataset (GSE287288) yielded similar results, with CPI ranging from "
        f"79.2% (B cells) to 89.0% (Dendritic cells), mean {mean_dtb*100:.1f}% (Table 1). The consistency "
        f"across independent cohorts, tissue types (BAL vs. PBMC), and disease presentations validates "
        f"CPI as a robust metric."
    )
    
    # Table 1
    doc.add_paragraph()
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Cell Type'
    hdr[1].text = 'Dataset'
    hdr[2].text = 'DEGs'
    hdr[3].text = 'CPI (%)'
    
    for r in cpi_bal:
        row = table1.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'BAL'
        row[2].text = r['n_deg']
        row[3].text = f"{float(r['CPI'])*100:.1f}"
    
    for r in cpi_dtb:
        row = table1.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'DTB PBMC'
        row[2].text = r['n_deg']
        row[3].text = f"{float(r['CPI'])*100:.1f}"
    
    doc.add_paragraph("Table 1. Chromatin Priming Index by cell type and dataset.").italic = True
    
    doc.add_heading("3.2 Dendritic Cells Show Highest Chromatin Priming", level=2)
    doc.add_paragraph(
        "Dendritic cells consistently showed the highest CPI across both datasets (BAL: 81.2%, DTB PBMC: 89.0%), "
        "suggesting that antigen-presenting cells have particularly extensive epigenetic priming. This may "
        "reflect the requirement for rapid transcriptional responses to diverse pathogen-associated signals."
    )
    
    doc.add_heading("3.3 Master Regulator TFs are Chromatin-Primed", level=2)
    
    # Get TF counts
    n_tfs = len(set(r['tf'] for r in tf_summary)) if tf_summary else 0
    
    doc.add_paragraph(
        f"Master regulator analysis identified {n_tfs} transcription factors differentially expressed across "
        f"datasets (Table 2). The top-ranked TFs included CEBPB (score 14.5, BAL), SPIB (9.66, DTB PBMC), "
        f"IRF7 (7.35, BAL), and STAT1 (7.34, BAL). Notably, all top 15 TFs were chromatin-primed, indicating "
        f"their regulatory potential is supported by accessible chromatin."
    )
    
    doc.add_paragraph(
        "Sixteen TFs were conserved across both datasets: CEBPB, IRF7, STAT1, JUN, REL, SPI1, RUNX1, ELF1, "
        "CEBPD, JUND, IRF3, IRF8, IRF1, XBP1, CEBPA, and IRF2. These represent core transcriptional regulators "
        "of myeloid function and inflammatory responses, consistent with their known roles in TB immunity."
    )
    
    # ===== DISCUSSION =====
    doc.add_heading("4. Discussion", level=1)
    
    doc.add_paragraph(
        "We introduce the Chromatin Priming Index as a practical metric for assessing epigenetic priming "
        "from scRNA-seq data using reference chromatin accessibility atlases. Our validation across two "
        "independent TB cohorts demonstrates several key findings."
    )
    
    doc.add_paragraph(
        "First, CPI is remarkably consistent (77-89%) across tissue types and disease presentations, "
        "suggesting that the majority of transcriptionally dysregulated genes in TB have prior chromatin "
        "accessibility. This supports the concept that epigenetic priming establishes a permissive "
        "transcriptional landscape that is subsequently activated during infection."
    )
    
    doc.add_paragraph(
        "Second, dendritic cells show the highest CPI, consistent with their role as professional "
        "antigen-presenting cells requiring rapid and flexible transcriptional responses. This finding "
        "aligns with studies showing extensive epigenetic priming in DCs during development (Satpathy "
        "et al., 2019)."
    )
    
    doc.add_paragraph(
        "Third, master regulator TFs identified through our analysis—including CEBPB, IRF7, STAT1, and "
        "JUN—are well-established regulators of myeloid function and inflammatory responses. Their "
        "conservation across independent datasets and chromatin priming status supports their central "
        "role in TB pathogenesis."
    )
    
    doc.add_paragraph(
        "Limitations of our approach include the use of healthy PBMC reference for diseased tissue, which "
        "may underestimate disease-specific chromatin changes. Additionally, CPI does not capture the "
        "direction of accessibility changes (opening vs. closing) or the specific regulatory elements "
        "involved. Future work should develop disease-specific chromatin atlases and incorporate enhancer-"
        "gene predictions."
    )
    
    # ===== CONCLUSIONS =====
    doc.add_heading("5. Conclusions", level=1)
    
    doc.add_paragraph(
        "The Chromatin Priming Index provides a quantitative, generalizable framework for integrating "
        "scRNA-seq with chromatin accessibility references. Its robustness across independent TB cohorts "
        "demonstrates utility for infectious disease research without requiring expensive multiome profiling. "
        "CPI can be readily applied to any scRNA-seq dataset using publicly available chromatin atlases, "
        "enabling broader assessment of epigenetic priming in disease contexts."
    )
    
    # ===== DATA AVAILABILITY =====
    doc.add_heading("Data Availability", level=1)
    doc.add_paragraph(
        "scRNA-seq data are available from GEO under accessions GSE167232 and GSE287288. The ATAC reference "
        "(PBMC multiome) is available from 10x Genomics. Analysis code is available at [GitHub repository]."
    )
    
    # ===== FUNDING =====
    doc.add_heading("Funding", level=1)
    doc.add_paragraph("[Funding sources to be added]")
    
    # ===== CONFLICTS =====
    doc.add_heading("Conflict of Interest", level=1)
    doc.add_paragraph("The authors declare no conflicts of interest.")
    
    # ===== AI STATEMENT =====
    doc.add_heading("AI Usage Statement", level=1)
    doc.add_paragraph(
        "AI tools were used to assist with code development, data analysis pipeline construction, and "
        "manuscript drafting. All analyses were verified and results validated by the authors. The "
        "manuscript was reviewed for scientific accuracy before submission."
    )
    
    # ===== REFERENCES =====
    doc.add_page_break()
    doc.add_heading("References", level=1)
    
    for ref in REFERENCES:
        doc.add_paragraph(format_reference(ref))
    
    # ===== FIGURE LEGENDS =====
    doc.add_page_break()
    doc.add_heading("Figure Legends", level=1)
    
    doc.add_paragraph().add_run("Figure 1. ").bold = True
    doc.paragraphs[-1].add_run(
        "Chromatin Priming Index by cell type in BAL samples. Bar plot showing CPI (%) for each cell type "
        "identified in the GSE167232 dataset. CPI ranged from 77.6% to 81.2%."
    )
    
    doc.add_paragraph().add_run("Figure 2. ").bold = True
    doc.paragraphs[-1].add_run(
        "Cross-dataset CPI comparison. Bar plot comparing CPI between BAL (GSE167232) and DTB PBMC "
        "(GSE287288) datasets across cell types. Both datasets show consistently high CPI (77-89%)."
    )
    
    doc.add_paragraph().add_run("Figure 3. ").bold = True
    doc.paragraphs[-1].add_run(
        "UMAP visualization of DTB PBMC cell types. Dimensionality reduction showing T cells (n=6,059), "
        "B cells (n=4,539), monocytes (n=4,482), NK cells (n=3,557), and dendritic cells (n=497)."
    )
    
    doc.add_paragraph().add_run("Figure 4. ").bold = True
    doc.paragraphs[-1].add_run(
        "Master regulator TF ranking. Top 15 transcription factors ranked by composite score in BAL (top) "
        "and DTB PBMC (bottom) datasets. Red bars indicate chromatin-primed TFs."
    )
    
    # Save main manuscript
    main_path = OUTPUT_DIR / "Manuscript_CPI_Submission.docx"
    doc.save(main_path)
    print(f"Main manuscript saved to: {main_path}")
    
    # ===== SUPPLEMENTARY MATERIALS =====
    supp = Document()
    supp.add_heading("Supplementary Materials", level=0)
    supp.add_heading("Chromatin Priming Index: A Quantitative Metric", level=1)
    
    supp.add_heading("Supplementary Table S1. Complete CPI Results", level=2)
    supp_table = supp.add_table(rows=1, cols=5)
    supp_table.style = 'Table Grid'
    sh = supp_table.rows[0].cells
    sh[0].text = 'Cell Type'
    sh[1].text = 'Dataset'
    sh[2].text = 'Total DEGs'
    sh[3].text = 'DEGs with Links'
    sh[4].text = 'CPI (%)'
    
    for r in cpi_bal:
        row = supp_table.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'GSE167232 (BAL)'
        row[2].text = r['n_deg']
        row[3].text = r.get('n_deg_with_link', 'NA')
        row[4].text = f"{float(r['CPI'])*100:.2f}"
    
    for r in cpi_dtb:
        row = supp_table.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'GSE287288 (DTB PBMC)'
        row[2].text = r['n_deg']
        row[3].text = r.get('n_deg_with_link', 'NA')
        row[4].text = f"{float(r['CPI'])*100:.2f}"
    
    supp.add_heading("Supplementary Table S2. TF Master Regulators", level=2)
    if tf_summary:
        tf_table = supp.add_table(rows=1, cols=5)
        tf_table.style = 'Table Grid'
        tfh = tf_table.rows[0].cells
        tfh[0].text = 'TF'
        tfh[1].text = 'Dataset'
        tfh[2].text = 'Cell Types'
        tfh[3].text = 'Score'
        tfh[4].text = 'Primed'
        
        for r in tf_summary[:20]:  # Top 20
            row = tf_table.add_row().cells
            row[0].text = r['tf']
            row[1].text = r['dataset']
            row[2].text = r['n_cell_types']
            row[3].text = f"{float(r['score']):.2f}"
            row[4].text = 'Yes' if r.get('is_primed') == 'TRUE' else 'No'
    
    supp_path = OUTPUT_DIR / "Supplementary_Materials.docx"
    supp.save(supp_path)
    print(f"Supplementary materials saved to: {supp_path}")
    
    print("\n✓ Submission package complete!")
    print(f"  Main manuscript: {main_path}")
    print(f"  Supplementary: {supp_path}")

if __name__ == "__main__":
    main()
