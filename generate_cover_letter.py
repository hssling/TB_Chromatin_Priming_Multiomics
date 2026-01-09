"""
Generate IJTLD Cover Letter
"""

from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "4_results" / "submission"

AUTHOR = {
    "name": "Dr. Siddalingaiah H S",
    "degree": "MD",
    "designation": "Professor",
    "department": "Department of Community Medicine",
    "institution": "Shridevi Institute of Medical Sciences and Research Hospital",
    "city": "Tumkur",
    "state": "Karnataka",
    "country": "India",
    "pin": "572106",
    "email": "hssling@yahoo.com",
    "phone": "+91-8941087719",
    "orcid": "0000-0002-4771-8285"
}

def main():
    print("Generating IJTLD cover letter...")
    
    doc = Document()
    
    # Date
    today = datetime.now().strftime("%d %B %Y")
    date_para = doc.add_paragraph(today)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Addressee
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("International Journal of Tuberculosis and Lung Disease")
    doc.add_paragraph("The Union")
    
    doc.add_paragraph()
    
    # Subject
    subject = doc.add_paragraph()
    subject.add_run("Subject: ").bold = True
    subject.add_run("Submission of Original Research Article")
    
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Editor,")
    
    doc.add_paragraph()
    
    # Body
    doc.add_paragraph(
        "I am pleased to submit our manuscript entitled \"Chromatin Priming Index: A Novel "
        "Quantitative Framework for Integrating Single-Cell Transcriptomic and Epigenomic Data "
        "in Tuberculosis\" for consideration as an Original Article in the International Journal "
        "of Tuberculosis and Lung Disease."
    )
    
    doc.add_paragraph(
        "Understanding the epigenetic regulation of immune responses in tuberculosis (TB) is "
        "critical for developing novel therapeutic strategies. However, matched multiome profiling "
        "remains expensive and technically challenging. In this manuscript, we introduce the "
        "Chromatin Priming Index (CPI), a practical metric that enables chromatin priming "
        "assessment from single-cell RNA sequencing data using publicly available chromatin "
        "accessibility references."
    )
    
    doc.add_paragraph(
        "We validated CPI using two independent TB cohorts: bronchoalveolar lavage macrophages "
        "(GSE167232, n=10,357 cells) and peripheral blood mononuclear cells from disseminated TB "
        "patients (GSE287288, n=21,000 cells), integrated with 1.28 million peak-gene links from "
        "10x Genomics PBMC multiome reference. Our key findings include:"
    )
    
    findings = [
        "CPI is consistently high (77-89%) across tissue types and cell populations",
        "Dendritic cells show the highest CPI (89%), suggesting extensive epigenetic priming",
        "16 master regulator transcription factors are conserved across datasets, including "
        "CEBPB, IRF7, STAT1, and JUN",
        "All top TFs are chromatin-primed, supporting their regulatory role in TB immunity"
    ]
    for f in findings:
        doc.add_paragraph(f"• {f}")
    
    doc.add_paragraph(
        "We believe this work makes a significant contribution to the field by providing a "
        "generalizable framework for integrating single-cell transcriptomic and epigenomic data "
        "without requiring expensive multiome profiling. The methodology can be readily applied "
        "to any scRNA-seq dataset using publicly available chromatin atlases."
    )
    
    doc.add_paragraph(
        "The manuscript contains approximately 3,000 words, 3 tables, 4 figures, and 12 references. "
        "Supplementary materials include extended methods, additional tables (S1-S4), and supplementary "
        "figures (S1-S4). All analysis code is available at our GitHub repository."
    )
    
    doc.add_paragraph(
        "This manuscript has not been published previously and is not under consideration for "
        "publication elsewhere. All authors have approved the manuscript and agree with its "
        "submission to IJTLD. There are no conflicts of interest to declare. This research received "
        "no specific funding."
    )
    
    doc.add_paragraph(
        "We confirm that AI-assisted tools were used for code development and manuscript preparation. "
        "All analyses were verified and the manuscript was reviewed for scientific accuracy by the author."
    )
    
    doc.add_paragraph(
        "Thank you for considering our manuscript. We look forward to your response."
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.add_run(AUTHOR['name']).bold = True
    doc.add_paragraph(f"{AUTHOR['degree']}, {AUTHOR['designation']}")
    doc.add_paragraph(AUTHOR['department'])
    doc.add_paragraph(AUTHOR['institution'])
    doc.add_paragraph(f"{AUTHOR['city']}, {AUTHOR['state']}, {AUTHOR['country']} - {AUTHOR['pin']}")
    doc.add_paragraph(f"Email: {AUTHOR['email']}")
    doc.add_paragraph(f"Phone: {AUTHOR['phone']}")
    doc.add_paragraph(f"ORCID: https://orcid.org/{AUTHOR['orcid']}")
    
    # Save
    output_path = OUTPUT_DIR / "CoverLetter_IJTLD.docx"
    doc.save(output_path)
    print(f"Cover letter saved to: {output_path}")

if __name__ == "__main__":
    main()
