"""
Generate submission-ready manuscript for Frontiers in Immunology
With verified author details and references
"""

import os
import csv
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "4_results"
TABLES_DIR = RESULTS_DIR / "tables"
FIGURES_DIR = RESULTS_DIR / "figures"
OUTPUT_DIR = RESULTS_DIR / "submission"
OUTPUT_DIR.mkdir(exist_ok=True)

# Author Details
AUTHOR = {
    "name": "Siddalingaiah H S",
    "degree": "MD",
    "designation": "Professor",
    "department": "Department of Community Medicine",
    "institution": "Shridevi Institute of Medical Sciences and Research Hospital",
    "city": "Tumkur",
    "state": "Karnataka",
    "country": "India",
    "pin": "572106",
    "email": "hssling@yahoo.com",
    "phone": "+91-8941087719",
    "orcid": "0000-0002-4771-8285"
}

# Verified references
REFERENCES = [
    "1. Pisu D, Huang L, Narang V, et al. Single cell analysis of M. tuberculosis phenotype and macrophage lineages in the infected lung. J Exp Med. 2021;218(9):e20210615. doi: 10.1084/jem.20210615",
    "2. Gong Z, Xu H, Zhang Q, Chen Y, Xie J. Unveiling the immunological landscape of disseminated tuberculosis: a single-cell transcriptome perspective. Front Immunol. 2025;16:1527592. doi: 10.3389/fimmu.2025.1527592",
    "3. Stuart T, Butler A, Hoffman P, et al. Comprehensive Integration of Single-Cell Data. Cell. 2019;177(7):1888-1902.e21. doi: 10.1016/j.cell.2019.05.031",
    "4. Hao Y, Hao S, Andersen-Nissen E, et al. Integrated analysis of multimodal single-cell data. Cell. 2021;184(13):3573-3587.e29. doi: 10.1016/j.cell.2021.04.048",
    "5. Netea MG, Dominguez-Andres J, Barreiro LB, et al. Defining trained immunity and its role in health and disease. Nat Rev Immunol. 2020;20(6):375-388. doi: 10.1038/s41577-020-0285-6",
    "6. Cheng SC, Quintin J, Cramer RA, et al. mTOR- and HIF-1alpha-mediated aerobic glycolysis as metabolic basis for trained immunity. Science. 2014;345(6204):1250684. doi: 10.1126/science.1250684",
    "7. Buenrostro JD, Wu B, Chang HY, Greenleaf WJ. ATAC-seq: A Method for Assaying Chromatin Accessibility Genome-Wide. Curr Protoc Mol Biol. 2015;109:21.29.1-21.29.9. doi: 10.1002/0471142727.mb2129s109",
    "8. Granja JM, Corces MR, Pierce SE, et al. ArchR is a scalable software package for integrative single-cell chromatin accessibility analysis. Nat Genet. 2021;53(3):403-411. doi: 10.1038/s41588-021-00790-6",
    "9. Satpathy AT, Granja JM, Yost KE, et al. Massively parallel single-cell chromatin landscapes of human immune cell development and intratumoral T cell exhaustion. Nat Biotechnol. 2019;37(8):925-936. doi: 10.1038/s41587-019-0206-z",
    "10. World Health Organization. Global tuberculosis report 2024. Geneva: WHO; 2024.",
]

def read_csv(path):
    if not path.exists():
        return []
    with open(path, 'r', encoding='utf-8') as f:
        return list(csv.DictReader(f))

def main():
    print("Generating Frontiers in Immunology manuscript...")
    
    cpi_bal = read_csv(TABLES_DIR / "Table_CPI_by_celltype.csv")
    cpi_dtb = read_csv(TABLES_DIR / "Table_CPI_GSE287288_DTB.csv")
    tf_summary = read_csv(TABLES_DIR / "Table_TF_master_regulators.csv")
    
    mean_bal = sum(float(r['CPI']) for r in cpi_bal) / len(cpi_bal) if cpi_bal else 0
    mean_dtb = sum(float(r['CPI']) for r in cpi_dtb) / len(cpi_dtb) if cpi_dtb else 0
    
    doc = Document()
    
    # TITLE
    title = doc.add_paragraph()
    title.add_run(
        "Chromatin Priming Index: A Quantitative Metric for Integrating "
        "Single-Cell Transcriptomic and Chromatin Accessibility Data in Tuberculosis"
    ).bold = True
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # AUTHOR
    author_para = doc.add_paragraph()
    author_para.add_run(f"{AUTHOR['name']}").bold = True
    author_para.add_run("*")
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # AFFILIATION
    affil = doc.add_paragraph()
    affil.add_run(f"{AUTHOR['department']}, {AUTHOR['institution']}, ")
    affil.add_run(f"{AUTHOR['city']}, {AUTHOR['state']}, {AUTHOR['country']}")
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # CORRESPONDENCE
    corresp = doc.add_paragraph()
    corresp.add_run("*Correspondence: ").bold = True
    corresp.add_run(f"{AUTHOR['name']}, {AUTHOR['email']}")
    corresp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ORCID
    orcid = doc.add_paragraph()
    orcid.add_run(f"ORCID: https://orcid.org/{AUTHOR['orcid']}")
    orcid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orcid.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # ABSTRACT
    doc.add_heading("Abstract", level=1)
    
    doc.add_paragraph().add_run("Background: ").bold = True
    doc.paragraphs[-1].add_run(
        "Understanding chromatin regulation in disease requires integration of transcriptomic and "
        "epigenomic data. However, matched single-cell multiome profiling is expensive and technically "
        "challenging, limiting chromatin priming assessment in most clinical cohorts."
    )
    
    doc.add_paragraph().add_run("Methods: ").bold = True
    doc.paragraphs[-1].add_run(
        f"We developed the Chromatin Priming Index (CPI), measuring the proportion of differentially "
        f"expressed genes with accessible chromatin regions in a reference atlas. Using 1,283,042 "
        f"peak-gene links from 10x Genomics PBMC multiome dataset, we validated CPI across two "
        f"independent tuberculosis cohorts: bronchoalveolar lavage macrophages (GSE167232, n=10,357 cells) "
        f"and peripheral blood from disseminated TB patients (GSE287288, n=21,000 cells)."
    )
    
    doc.add_paragraph().add_run("Results: ").bold = True
    doc.paragraphs[-1].add_run(
        f"CPI was consistently high across tissue types (BAL: {mean_bal*100:.1f}%, PBMC: {mean_dtb*100:.1f}%) "
        f"and cell types (range: 77.6-89.0%). Dendritic cells exhibited the highest CPI (89.0%). "
        f"Master regulator analysis identified 16 transcription factors conserved across datasets, "
        f"including CEBPB, IRF7, STAT1, and JUN, all showing chromatin accessibility support."
    )
    
    doc.add_paragraph().add_run("Conclusions: ").bold = True
    doc.paragraphs[-1].add_run(
        "CPI provides a quantitative framework for assessing chromatin priming from scRNA-seq data "
        "without requiring matched ATAC-seq profiling. Its consistency across independent TB cohorts "
        "demonstrates utility for infectious disease research."
    )
    
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("chromatin priming, single-cell RNA-seq, ATAC-seq, data integration, tuberculosis, transcription factors")
    
    doc.add_page_break()
    
    # INTRODUCTION
    doc.add_heading("Introduction", level=1)
    
    doc.add_paragraph(
        "Tuberculosis (TB) remains a leading cause of mortality from a single infectious agent, "
        "with an estimated 10.6 million new cases and 1.3 million deaths globally in 2022 (10). "
        "Single-cell RNA sequencing (scRNA-seq) has transformed our understanding of cellular "
        "heterogeneity in TB (1, 2), but understanding the regulatory mechanisms driving these "
        "transcriptional states requires knowledge of chromatin accessibility (7)."
    )
    
    doc.add_paragraph(
        "Chromatin priming, whereby accessible chromatin regions precede or enable transcriptional "
        "activation, is central to rapid immune responses and trained immunity (5). In TB, epigenetic "
        "reprogramming of myeloid cells contributes to both protective immunity and disease pathology (6). "
        "However, direct assessment of chromatin priming typically requires matched multiome profiling, "
        "which remains expensive and technically demanding (8)."
    )
    
    doc.add_paragraph(
        "We propose the Chromatin Priming Index (CPI) as a practical metric that leverages publicly "
        "available chromatin accessibility references to assess epigenetic priming from scRNA-seq data "
        "alone. Here, we validate CPI using two independent TB cohorts and demonstrate its utility for "
        "identifying master regulator transcription factors with chromatin accessibility support."
    )
    
    # MATERIALS AND METHODS
    doc.add_heading("Materials and Methods", level=1)
    
    doc.add_heading("CPI Definition", level=2)
    doc.add_paragraph(
        "The Chromatin Priming Index was calculated as: CPI = |DEGs with peak-gene links| / |Total DEGs|, "
        "where DEGs are statistically significant differentially expressed genes (adjusted P < 0.05) and "
        "peak-gene links represent associations between accessible chromatin peaks and genes derived from "
        "correlation analysis in multiome reference data. CPI ranges from 0 to 1."
    )
    
    doc.add_heading("Reference Atlas", level=2)
    doc.add_paragraph(
        "Peak-gene links were obtained from the 10x Genomics PBMC multiome dataset (12,012 cells, "
        "111,857 peaks). Feature linkage analysis identified 1,283,042 peak-gene associations based on "
        "Pearson correlation between peak accessibility and gene expression across cells."
    )
    
    doc.add_heading("Datasets", level=2)
    doc.add_paragraph(
        "Two independent TB scRNA-seq datasets were analyzed: (i) GSE167232: Bronchoalveolar lavage (BAL) "
        "samples containing 10,357 cells from Pisu et al. (1); (ii) GSE287288: PBMCs from seven patients "
        "with disseminated TB containing 21,000 cells from Gong et al. (2)."
    )
    
    doc.add_heading("Analysis Pipeline", level=2)
    doc.add_paragraph(
        "scRNA-seq data processing was performed using Seurat v5 (4). Standard preprocessing included "
        "quality control, normalization, identification of 3,000 highly variable genes, scaling, PCA, "
        "and UMAP visualization. Cell types were annotated using canonical marker gene expression. "
        "Differential expression analysis was performed using the Wilcoxon rank-sum test."
    )
    
    doc.add_heading("Master Regulator Analysis", level=2)
    doc.add_paragraph(
        "Transcription factors were identified from a curated list of 67 immune-relevant TFs spanning "
        "NFkB, AP-1, IRF, STAT, and C/EBP families. TFs differentially expressed in multiple cell types "
        "were ranked by a composite score incorporating expression magnitude, statistical significance, "
        "and chromatin priming status."
    )
    
    doc.add_heading("Statistical Analysis", level=2)
    doc.add_paragraph(
        "Fisher's exact test was used for enrichment analyses. Wilcoxon rank-sum tests were used for "
        "comparing accessibility scores. All analyses were performed in R 4.5.1 with significance "
        "threshold alpha = 0.05."
    )
    
    # RESULTS
    doc.add_heading("Results", level=1)
    
    doc.add_heading("CPI is Consistent Across Tissue Types", level=2)
    doc.add_paragraph(
        f"We applied CPI to the BAL dataset (GSE167232). Across five cell types, CPI ranged from "
        f"77.6% (Interstitial Macrophages) to 81.2% (Dendritic cells), with a mean of {mean_bal*100:.1f}% "
        f"(Table 1). Application to the DTB PBMC dataset (GSE287288) yielded similar results, with CPI "
        f"ranging from 79.2% (B cells) to 89.0% (Dendritic cells), mean {mean_dtb*100:.1f}%."
    )
    
    # Table 1
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Cell Type'
    hdr[1].text = 'Dataset'
    hdr[2].text = 'DEGs'
    hdr[3].text = 'CPI (%)'
    
    for r in cpi_bal:
        row = table1.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'BAL'
        row[2].text = r['n_deg']
        row[3].text = f"{float(r['CPI'])*100:.1f}"
    
    for r in cpi_dtb:
        row = table1.add_row().cells
        row[0].text = r['cell_type'].replace('_', ' ')
        row[1].text = 'DTB PBMC'
        row[2].text = r['n_deg']
        row[3].text = f"{float(r['CPI'])*100:.1f}"
    
    doc.add_paragraph("Table 1. Chromatin Priming Index by cell type and dataset.").italic = True
    
    doc.add_heading("Dendritic Cells Show Highest Chromatin Priming", level=2)
    doc.add_paragraph(
        "Dendritic cells consistently showed the highest CPI across both datasets (BAL: 81.2%, "
        "DTB PBMC: 89.0%), suggesting that antigen-presenting cells have particularly extensive "
        "epigenetic priming for rapid transcriptional responses."
    )
    
    doc.add_heading("Master Regulator TFs are Chromatin-Primed", level=2)
    n_tfs = len(set(r['tf'] for r in tf_summary)) if tf_summary else 0
    doc.add_paragraph(
        f"Master regulator analysis identified {n_tfs} transcription factors differentially expressed "
        f"across datasets. Top-ranked TFs included CEBPB (score 14.5), SPIB (9.66), IRF7 (7.35), and "
        f"STAT1 (7.34). Notably, all top 15 TFs were chromatin-primed. Sixteen TFs were conserved "
        f"across both datasets: CEBPB, IRF7, STAT1, JUN, REL, SPI1, RUNX1, ELF1, CEBPD, JUND, IRF3, "
        f"IRF8, IRF1, XBP1, CEBPA, and IRF2."
    )
    
    # DISCUSSION
    doc.add_heading("Discussion", level=1)
    
    doc.add_paragraph(
        "We introduce the Chromatin Priming Index as a practical metric for assessing epigenetic "
        "priming from scRNA-seq data using reference chromatin accessibility atlases. Our validation "
        "across two independent TB cohorts demonstrates that CPI is remarkably consistent (77-89%) "
        "across tissue types and disease presentations."
    )
    
    doc.add_paragraph(
        "Dendritic cells showed the highest CPI, consistent with their role as professional "
        "antigen-presenting cells requiring rapid transcriptional responses. This aligns with studies "
        "showing extensive epigenetic priming in DCs during development (9). Master regulator TFs "
        "identified through our analysis, including CEBPB, IRF7, and STAT1, are well-established "
        "regulators of myeloid function and inflammatory responses."
    )
    
    doc.add_paragraph(
        "Limitations include the use of healthy PBMC reference for diseased tissue, which may "
        "underestimate disease-specific chromatin changes. Future work should develop disease-specific "
        "chromatin atlases and incorporate enhancer-gene predictions."
    )
    
    doc.add_paragraph(
        "In conclusion, CPI provides a quantitative framework for integrating scRNA-seq with chromatin "
        "accessibility references without requiring expensive multiome profiling. Its robustness across "
        "independent TB cohorts demonstrates utility for infectious disease research."
    )
    
    # DATA AVAILABILITY
    doc.add_heading("Data Availability Statement", level=1)
    doc.add_paragraph(
        "Publicly available datasets were analyzed in this study. scRNA-seq data are available from "
        "GEO under accessions GSE167232 and GSE287288. The ATAC reference (PBMC multiome) is available "
        "from 10x Genomics."
    )
    
    # ETHICS
    doc.add_heading("Ethics Statement", level=1)
    doc.add_paragraph(
        "This study analyzed publicly available datasets. No new human samples were collected. "
        "The original studies obtained appropriate ethical approvals."
    )
    
    # AUTHOR CONTRIBUTIONS
    doc.add_heading("Author Contributions", level=1)
    doc.add_paragraph(
        f"{AUTHOR['name']}: Conceptualization, Methodology, Software, Formal analysis, "
        f"Writing - original draft, Writing - review & editing, Visualization."
    )
    
    # FUNDING
    doc.add_heading("Funding", level=1)
    doc.add_paragraph(
        "This research received no specific grant from any funding agency in the public, commercial, "
        "or not-for-profit sectors."
    )
    
    # CONFLICT OF INTEREST
    doc.add_heading("Conflict of Interest", level=1)
    doc.add_paragraph("The author declares that the research was conducted in the absence of any "
                      "commercial or financial relationships that could be construed as a potential "
                      "conflict of interest.")
    
    # ACKNOWLEDGMENTS
    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph(
        "The author thanks the data generators of GSE167232 and GSE287288 for making their data "
        "publicly available. AI-assisted tools were used for code development and manuscript "
        "preparation; all analyses were verified and the manuscript was reviewed for scientific "
        "accuracy by the author."
    )
    
    # REFERENCES
    doc.add_page_break()
    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref)
    
    # FIGURE LEGENDS
    doc.add_page_break()
    doc.add_heading("Figure Legends", level=1)
    
    doc.add_paragraph().add_run("Figure 1. ").bold = True
    doc.paragraphs[-1].add_run("Chromatin Priming Index by cell type. CPI (%) for each cell type in BAL (GSE167232) and DTB PBMC (GSE287288) datasets.")
    
    doc.add_paragraph().add_run("Figure 2. ").bold = True
    doc.paragraphs[-1].add_run("UMAP visualization of DTB PBMC cell types from GSE287288.")
    
    doc.add_paragraph().add_run("Figure 3. ").bold = True
    doc.paragraphs[-1].add_run("Master regulator TF ranking by composite score. Red bars indicate chromatin-primed TFs.")
    
    # Save
    output_path = OUTPUT_DIR / "Manuscript_Frontiers_Immunology_FINAL.docx"
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")

if __name__ == "__main__":
    main()
